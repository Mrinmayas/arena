"""Render an operator-facing Word guide (.docx) from a Process.

Same source of truth as the README (``PROCESS_STEPS.md`` → :class:`core.docs.Process`),
but formatted for a non-developer operator who prefers Word: each step is grouped by
phase and tagged **[Automated]** (the tool does it) or **[You]** (a person does it),
human-review (HITL) steps are called out, decisions are listed as branches, and a
screenshot is embedded at any step that names one.

Requires the ``docs`` extra: ``uv sync --extra docs`` (python-docx).
"""
from __future__ import annotations

from pathlib import Path

from .steps import Process


def render_operator_guide(
    process: Process,
    out_path: str | Path,
    *,
    screenshots_dir: str | Path | None = None,
) -> Path:
    """Write an operator .docx guide for ``process``; returns the output path.

    ``screenshots_dir`` (optional): a folder of ``<screenshot-id>.png`` images; a step's
    ``Screenshot:`` id is embedded when the matching file exists.
    """
    from docx import Document  # local import: only needed when actually rendering
    from docx.shared import Inches

    doc = Document()
    doc.add_heading(f"{process.name} — Operator Guide", level=0)
    if process.context:
        doc.add_paragraph(process.context, style="Intense Quote")

    doc.add_heading("How to use this", level=1)
    doc.add_paragraph(
        "This guide walks the process step by step. Steps marked [Automated] are done "
        "for you by the tool; steps marked [You] are actions you perform. Steps that "
        "need your review before continuing are flagged “Review needed”."
    )

    shots = Path(screenshots_dir) if screenshots_dir else None

    for phase, steps in process.by_phase().items():
        if phase:
            doc.add_heading(phase, level=1)
        for s in steps:
            doc.add_heading(f"Step {s.number} — {s.title}", level=2)

            who = "Automated" if s.is_automated else "You"
            tags = [f"[{who}]"]
            if s.hitl:
                tags.append("Review needed")
            tag_p = doc.add_paragraph()
            run = tag_p.add_run(" · ".join(tags))
            run.bold = True

            if s.description:
                doc.add_paragraph(s.description)

            if s.decision is not None:
                dp = doc.add_paragraph()
                dp.add_run("Decision: ").bold = True
                dp.add_run(s.decision.question)
                for b in s.decision.branches:
                    doc.add_paragraph(f"{b.label} → {b.target}", style="List Bullet")

            if s.note:
                np = doc.add_paragraph()
                np.add_run("Note: ").bold = True
                np.add_run(s.note)

            if s.screenshot and shots is not None:
                img = shots / f"{s.screenshot}.png"
                if img.exists():
                    doc.add_picture(str(img), width=Inches(6.0))

    out = Path(out_path)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    return out
